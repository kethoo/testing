import streamlit as st
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
import re, io

def highlight_in_runs(runs, patterns):
    """Highlight matches inside a list of runs (paragraph or table cell)."""
    text = "".join(run.text for run in runs)
    for pattern in patterns:
        if re.search(pattern, text):
            # Rebuild runs with highlighting
            new_runs = []
            last_end = 0
            for match in re.finditer(pattern, text):
                if match.start() > last_end:
                    new_runs.append((text[last_end:match.start()], False))
                new_runs.append((match.group(), True))
                last_end = match.end()
            if last_end < len(text):
                new_runs.append((text[last_end:], False))

            # Clear old runs
            for run in runs:
                run.text = ""
            # Add new runs
            for run_text, is_highlight in new_runs:
                run = runs[0].paragraph.add_run(run_text)
                if is_highlight:
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            break  # no need to re-check with the same pattern

def highlight_docx(file, terms):
    doc = Document(file)
    patterns = [re.compile(rf"({re.escape(term)})", re.IGNORECASE) for term in terms]

    # Highlight in paragraphs
    for para in doc.paragraphs:
        highlight_in_runs(para.runs, patterns)

    # Highlight in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    highlight_in_runs(para.runs, patterns)

    output = io.BytesIO()
    doc.save(output)
    return output

# --- Streamlit App ---
st.title("📄 Word Highlighter")
st.write("Upload a Word file, enter words, and download a highlighted copy.")

file = st.file_uploader("Upload a Word document", type="docx")
terms = st.text_input("Enter words to highlight (comma-separated)")

if file and terms:
    words = [t.strip() for t in terms.split(",") if t.strip()]
    if words:
        highlighted = highlight_docx(file, words)
        st.download_button(
            "⬇️ Download Highlighted File",
            highlighted.getvalue(),
            file_name="highlighted.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
